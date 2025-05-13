from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from sentence_transformers import SentenceTransformer, util
from PyPDF2 import PdfReader
import docx
import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image
import re
import nltk
from nltk.tokenize import sent_tokenize
import io
import traceback

nltk.download('punkt')  # Tokenizer

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Update for production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

model = SentenceTransformer("all-MiniLM-L6-v2")

def extract_text_from_pdf(contents: bytes) -> str:
    text = ""
    try:
        reader = PdfReader(io.BytesIO(contents))
        for page in reader.pages:
            page_text = page.extract_text() or ""
            text += page_text
    except Exception as e:
        print(f"PDF extraction error: {e}")
    if not text.strip():
        try:
            images = convert_from_bytes(contents)
            for image in images:
                text += pytesseract.image_to_string(image)
        except Exception as e:
            print(f"PDF OCR error: {e}")
    return text

def extract_text_from_docx(contents: bytes) -> str:
    try:
        doc = docx.Document(io.BytesIO(contents))
        text = "\n".join([para.text for para in doc.paragraphs])
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + " | ".join(cell.text for cell in row.cells)
        return text
    except Exception as e:
        print(f"DOCX extraction error: {e}")
        return ""

def get_text_from_file(filename: str, contents: bytes) -> str:
    ext = filename.lower().split('.')[-1]
    if ext == "pdf":
        return extract_text_from_pdf(contents)
    elif ext == "docx":
        return extract_text_from_docx(contents)
    elif ext in ("txt", "json"):
        return contents.decode("utf-8")
    else:
        raise ValueError(f"Unsupported file format: {ext}")

def extract_questions(text: str):
    lines = text.split('\n')
    questions = []
    for line in lines:
        line = line.strip()
        if not line or len(line) < 10:
            continue
        if re.search(r'\b(who|what|when|where|why|how|define|explain|describe|list|mention)\b', line.lower()):
            if not re.search(r'\bmarks?\b|\bsection\b|\bpart\b|^\d+\.?$|^time\b|^\(?[0-9]+\)?$', line.lower()):
                questions.append(line)
    return questions

def compute_semantic_similarity(textA: str, textB: str) -> float:
    sentsA = [s.strip() for s in sent_tokenize(textA) if len(s.split()) > 3]
    sentsB = [s.strip() for s in sent_tokenize(textB) if len(s.split()) > 3]

    if not sentsA or not sentsB:
        return 0.0

    embeddingsA = model.encode(sentsA, convert_to_tensor=True)
    embeddingsB = model.encode(sentsB, convert_to_tensor=True)

    sim_matrix = util.pytorch_cos_sim(embeddingsA, embeddingsB)
    max_sims = sim_matrix.max(dim=1).values
    avg_sim = max_sims.mean().item()

    return round(avg_sim * 100, 2)

@app.post("/compare")
async def compare(fileA: UploadFile = File(...), fileB: UploadFile = File(...)):
    try:
        print(f"📁 Received fileA: {fileA.filename}, fileB: {fileB.filename}")
        contentsA = await fileA.read()
        contentsB = await fileB.read()

        textA = get_text_from_file(fileA.filename, contentsA)
        textB = get_text_from_file(fileB.filename, contentsB)

        print(f"✅ Text A length: {len(textA)}")
        print(f"✅ Text B length: {len(textB)}")

        questionsA = extract_questions(textA)
        questionsB = extract_questions(textB)

        print(f"🔍 Extracted {len(questionsA)} questions from fileA")
        print(f"🔍 Extracted {len(questionsB)} questions from fileB")

        if len(questionsA) < 2 or len(questionsB) < 2:
            raise HTTPException(status_code=400, detail="Not enough questions in one or both files.")

        similarity = compute_semantic_similarity(" ".join(questionsA), " ".join(questionsB))

        return {"similarity": similarity}

    except ValueError as e:
        raise HTTPException(status_code=400, detail=f"Unsupported file format: {str(e)}")
    except Exception as e:
        print("❌ Internal Exception:")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Internal error: {str(e)}")
